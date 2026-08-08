from __future__ import annotations

import argparse
import copy
import hashlib
import json
import re
import tempfile
import threading
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw


def stable_hash(value: Any) -> str:
    payload = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def iter_block_items(parent: DocumentType | _Cell) -> Iterable[Paragraph | Table]:
    parent_element = parent.element.body if isinstance(parent, DocumentType) else parent._tc
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def paragraph_kind(paragraph: Paragraph) -> tuple[str, int | None]:
    style = paragraph.style.name if paragraph.style else ""
    if style.startswith("Heading "):
        try:
            return "heading", int(style.split(" ", 1)[1])
        except ValueError:
            return "heading", None
    if style.startswith("List"):
        return "list", None
    return "paragraph", None


def image_relationship_ids(paragraph: Paragraph) -> list[str]:
    return list(paragraph._p.xpath(".//a:blip/@r:embed"))


def extract_canonical(document: DocumentType) -> dict[str, Any]:
    assets: dict[str, bytes] = {}
    nodes: list[dict[str, Any]] = []
    order = 0

    for block in iter_block_items(document):
        order += 1
        if isinstance(block, Paragraph):
            kind, level = paragraph_kind(block)
            image_ids: list[str] = []
            for relationship_id in image_relationship_ids(block):
                part = document.part.related_parts[relationship_id]
                asset_id = hashlib.sha256(part.blob).hexdigest()[:16]
                assets[asset_id] = part.blob
                image_ids.append(asset_id)
            nodes.append(
                {
                    "nodeId": f"node-{order:03d}",
                    "nodeType": kind,
                    "order": order,
                    "text": block.text,
                    "headingLevel": level,
                    "style": block.style.name if block.style else None,
                    "imageAssetIds": image_ids,
                }
            )
        else:
            nodes.append(
                {
                    "nodeId": f"node-{order:03d}",
                    "nodeType": "table",
                    "order": order,
                    "rows": [[cell.text for cell in row.cells] for row in block.rows],
                }
            )

    return {
        "schemaVersion": "grant-docx-spike-v1",
        "nodes": nodes,
        "assets": {asset_id: blob.hex() for asset_id, blob in sorted(assets.items())},
    }


def configure_styles(document: DocumentType) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size in ((1, 16), (2, 13), (3, 11)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True


def create_docx_fixture(output_path: Path, image_path: Path) -> None:
    image = Image.new("RGB", (720, 300), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((40, 80, 210, 220), radius=16, fill="#E8F0FE", outline="#1A73E8", width=4)
    draw.rounded_rectangle((275, 80, 445, 220), radius=16, fill="#E6F4EA", outline="#188038", width=4)
    draw.rounded_rectangle((510, 80, 680, 220), radius=16, fill="#FEF7E0", outline="#F9AB00", width=4)
    draw.line((210, 150, 275, 150), fill="#5F6368", width=5)
    draw.line((445, 150, 510, 150), fill="#5F6368", width=5)
    image.save(image_path)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.text = "国家自然科学基金项目申请书（Spike 样例）"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer.paragraphs[0]
    footer.text = "ResearchGPT 技术验证"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("跨尺度生物信息整合的机制研究")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(20)

    document.add_heading("1 立项依据", level=1)
    document.add_heading("1.1 研究背景", level=2)
    document.add_paragraph(
        "复杂生命过程由多尺度信号共同调控，但现有研究常将结构、动力学与功能割裂分析。"
        "本项目拟建立可验证的跨尺度关联框架，并以申请人前期数据作为边界条件。"
    )
    document.add_paragraph("明确关键科学问题与可检验假说。", style="List Bullet")
    document.add_paragraph("建立结构—动力学—功能的证据链。", style="List Bullet")

    document.add_heading("1.2 研究现状与缺口", level=2)
    document.add_paragraph(
        "已有研究支持局部调控效应[1]，但尚未解释不同时间尺度之间的因果传递。"
        "现有结论在样本类型和实验边界上存在明显差异[2]。"
    )
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(5.6))
    caption = document.add_paragraph("图 1 研究假说的三阶段验证框架")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("2 研究目标", level=1)
    document.add_paragraph("本项目设置三个相互约束的研究目标，并通过统一评价指标验证。")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.2), Inches(3.7), Inches(1.4)]
    for index, value in enumerate(("目标", "可检验内容", "输出")):
        cell = table.rows[0].cells[index]
        cell.text = value
        cell.width = widths[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_values in (
        ("目标一", "识别跨尺度耦合的关键变量", "变量集合"),
        ("目标二", "验证耦合关系的因果方向", "机制模型"),
        ("目标三", "评估模型的适用边界", "验证报告"),
    ):
        row = table.add_row()
        for index, value in enumerate(row_values):
            row.cells[index].text = value
            row.cells[index].width = widths[index]
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_heading("参考文献", level=1)
    document.add_paragraph("[1] Zhang A, et al. Example mechanistic study. Journal A, 2024.")
    document.add_paragraph("[2] Li B, et al. Boundary conditions in multiscale systems. Journal B, 2025.")

    # A second section makes section/header fidelity observable during no-op export.
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("附录：验证清单", level=1)
    document.add_paragraph("本页用于观察分页、页眉页脚和分节属性能否稳定往返。")
    document.save(output_path)


def export_canonical(canonical: dict[str, Any], output_path: Path, asset_dir: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    configure_styles(document)

    assets: dict[str, str] = canonical["assets"]
    for node in canonical["nodes"]:
        kind = node["nodeType"]
        if kind == "heading":
            document.add_heading(node["text"], level=node.get("headingLevel") or 1)
        elif kind == "list":
            document.add_paragraph(node["text"], style=node.get("style") or "List Bullet")
        elif kind == "paragraph":
            paragraph = document.add_paragraph(node["text"])
            for asset_id in node.get("imageAssetIds", []):
                asset_path = asset_dir / f"{asset_id}.png"
                asset_path.write_bytes(bytes.fromhex(assets[asset_id]))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(asset_path), width=Inches(5.6))
        elif kind == "table":
            rows = node["rows"]
            table = document.add_table(rows=0, cols=len(rows[0]))
            table.style = "Table Grid"
            for values in rows:
                row = table.add_row()
                for index, value in enumerate(values):
                    row.cells[index].text = value
                    row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.save(output_path)


def canonical_metrics(canonical: dict[str, Any]) -> dict[str, Any]:
    nodes = canonical["nodes"]
    return {
        "nodeCount": len(nodes),
        "paragraphCount": sum(node["nodeType"] == "paragraph" for node in nodes),
        "headingCount": sum(node["nodeType"] == "heading" for node in nodes),
        "listCount": sum(node["nodeType"] == "list" for node in nodes),
        "tableCount": sum(node["nodeType"] == "table" for node in nodes),
        "imageCount": sum(len(node.get("imageAssetIds", [])) for node in nodes),
        "textHash": stable_hash(
            [node.get("text", node.get("rows")) for node in nodes]
        ),
        "assetHashes": sorted(canonical["assets"].keys()),
    }


def package_metrics(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        document_xml = archive.read("word/document.xml").decode("utf-8")
        return {
            "sectionCount": document_xml.count("<w:sectPr"),
            "headerPartCount": sum(name.startswith("word/header") and name.endswith(".xml") for name in names),
            "footerPartCount": sum(name.startswith("word/footer") and name.endswith(".xml") for name in names),
            "mediaPartCount": sum(name.startswith("word/media/") and not name.endswith("/") for name in names),
            "hasNumberingPart": "word/numbering.xml" in names,
            "hasStylesPart": "word/styles.xml" in names,
        }


def run_docx_roundtrip(output_dir: Path) -> dict[str, Any]:
    spike_dir = output_dir / "docx-roundtrip"
    spike_dir.mkdir(parents=True, exist_ok=True)
    source_path = spike_dir / "source.docx"
    exported_path = spike_dir / "roundtrip.docx"
    image_path = spike_dir / "fixture-image.png"
    create_docx_fixture(source_path, image_path)

    source_document = Document(source_path)
    source_canonical = extract_canonical(source_document)
    export_canonical(source_canonical, exported_path, spike_dir)
    exported_canonical = extract_canonical(Document(exported_path))

    source_metrics = canonical_metrics(source_canonical)
    exported_metrics = canonical_metrics(exported_canonical)
    body_equal = source_metrics == exported_metrics
    limitations = [
        "页眉、页脚和分节属性没有进入最小 Canonical Node 合同，no-op 导出无法保真。",
        "直接字体、段落间距、表格列宽和部分列表编号属于样式/布局层，当前 Spike 只验证结构，不承诺任意 DOCX 样式往返。",
        "浮动形状、文本框、域代码、批注、修订、脚注和复杂公式未纳入该最小样例，导入时必须生成 fidelity warning，不能静默丢失。",
    ]
    report = {
        "spike": "docx-roundtrip",
        "source": str(source_path),
        "roundtrip": str(exported_path),
        "sourceMetrics": source_metrics,
        "roundtripMetrics": exported_metrics,
        "bodyStructurePreserved": body_equal,
        "sourcePackageMetrics": package_metrics(source_path),
        "roundtripPackageMetrics": package_metrics(exported_path),
        "knownLimitations": limitations,
        "visualQa": {"status": "pending_external_render"},
    }
    write_json(spike_dir / "report.json", report)
    return report


@dataclass(frozen=True)
class Anchor:
    section_role: str
    heading: str
    text: str
    previous_text: str
    next_text: str
    node_type: str = "paragraph"


def normalize_text(text: str) -> str:
    return re.sub(r"[^\w\u4e00-\u9fff]+", "", text.lower())


def ngrams(text: str, size: int = 2) -> set[str]:
    normalized = normalize_text(text)
    if len(normalized) < size:
        return {normalized} if normalized else set()
    return {normalized[index : index + size] for index in range(len(normalized) - size + 1)}


def overlap(left: str, right: str) -> float:
    left_parts, right_parts = ngrams(left), ngrams(right)
    if not left_parts or not right_parts:
        return 0.0
    return len(left_parts & right_parts) / len(left_parts | right_parts)


def anchor_score(anchor: Anchor, candidate: dict[str, str]) -> float:
    score = 0.60 * overlap(anchor.text, candidate["text"])
    score += 0.12 * overlap(anchor.heading, candidate["heading"])
    score += 0.12 * (anchor.section_role == candidate["sectionRole"])
    score += 0.08 * overlap(anchor.previous_text, candidate.get("previousText", ""))
    score += 0.08 * overlap(anchor.next_text, candidate.get("nextText", ""))
    if anchor.node_type != candidate.get("nodeType", "paragraph"):
        score -= 0.20
    return round(max(score, 0.0), 4)


def locate_anchor(anchor: Anchor, candidates: list[dict[str, str]]) -> dict[str, Any]:
    ranked = sorted(
        ({"nodeId": candidate["nodeId"], "score": anchor_score(anchor, candidate)} for candidate in candidates),
        key=lambda item: item["score"],
        reverse=True,
    )
    best = ranked[0] if ranked else {"nodeId": None, "score": 0.0}
    runner_up = ranked[1]["score"] if len(ranked) > 1 else 0.0
    margin = round(best["score"] - runner_up, 4)
    if best["score"] >= 0.82 and margin >= 0.12:
        status = "matched"
    elif best["score"] >= 0.58:
        status = "ambiguous"
    else:
        status = "not_found"
    return {"status": status, "best": best, "margin": margin, "ranked": ranked[:3]}


def candidate(node_id: str, text: str, heading: str = "1.1 研究背景", role: str = "background", previous: str = "", next_text: str = "", node_type: str = "paragraph") -> dict[str, str]:
    return {
        "nodeId": node_id,
        "text": text,
        "heading": heading,
        "sectionRole": role,
        "previousText": previous,
        "nextText": next_text,
        "nodeType": node_type,
    }


def run_anchor_drift(output_dir: Path) -> dict[str, Any]:
    anchor = Anchor(
        section_role="background",
        heading="1.1 研究背景",
        text="现有研究尚未解释不同时间尺度之间的因果传递。",
        previous_text="复杂生命过程由多尺度信号共同调控。",
        next_text="本项目拟建立可验证的跨尺度关联框架。",
    )
    distractor = candidate("noise", "现有研究已经形成统一的评价框架。", heading="1.2 国内外现状", role="state_of_art")
    cases = {
        "insert_unrelated": [distractor, candidate("target", anchor.text, previous=anchor.previous_text, next_text=anchor.next_text)],
        "delete_unrelated": [candidate("target", anchor.text, previous=anchor.previous_text, next_text=anchor.next_text)],
        "split_paragraph": [candidate("target-a", "现有研究尚未解释不同时间尺度之间", previous=anchor.previous_text), candidate("target-b", "的因果传递。", previous="现有研究尚未解释不同时间尺度之间", next_text=anchor.next_text)],
        "merge_paragraph": [candidate("target", anchor.previous_text + anchor.text + anchor.next_text, previous="", next_text="")],
        "move_section": [candidate("target", anchor.text, heading="3.2 关键科学问题", role="scientific_question", previous=anchor.previous_text, next_text=anchor.next_text)],
        "rename_heading": [candidate("target", anchor.text, heading="1.1 研究基础与问题", previous=anchor.previous_text, next_text=anchor.next_text)],
        "paraphrase": [candidate("target", "不同观测尺度间的因果链条仍缺乏机制解释。", previous=anchor.previous_text, next_text=anchor.next_text)],
        "table_change": [candidate("table-target", anchor.text, previous=anchor.previous_text, next_text=anchor.next_text, node_type="table")],
    }
    results = {name: locate_anchor(anchor, values) for name, values in cases.items()}
    expected = {
        "insert_unrelated": "auto_match",
        "delete_unrelated": "auto_match",
        "split_paragraph": "human_review",
        "merge_paragraph": "human_review",
        "move_section": "human_review",
        "rename_heading": "auto_match",
        "paraphrase": "human_review",
        "table_change": "human_review",
    }
    passed = {
        name: (results[name]["status"] == "matched") if outcome == "auto_match" else (results[name]["status"] != "matched")
        for name, outcome in expected.items()
    }
    report = {
        "spike": "anchor-drift",
        "cases": results,
        "expectedClassification": expected,
        "safeDecisionPassRate": sum(passed.values()) / len(passed),
        "automaticMatchRate": sum(value["status"] == "matched" for value in results.values()) / len(results),
        "decision": "只对高分且有明显候选间隔的结果自动重定位；拆分、合并、跨语义章节移动、改写和节点类型变化必须进入人工确认。",
    }
    write_json(output_dir / "anchor-drift" / "report.json", report)
    return report


class RevisionStore:
    def __init__(self, nodes: dict[str, str]) -> None:
        self._lock = threading.Lock()
        self.current_revision = 1
        self.nodes = copy.deepcopy(nodes)
        self.audit: list[dict[str, Any]] = []

    def commit(self, *, base_revision: int, node_id: str, expected_hash: str, new_text: str, actor: str) -> dict[str, Any]:
        with self._lock:
            if base_revision != self.current_revision:
                return {"status": "revision_conflict", "currentRevision": self.current_revision}
            current_text = self.nodes[node_id]
            if stable_hash(current_text) != expected_hash:
                return {"status": "content_hash_conflict", "currentRevision": self.current_revision}
            self.nodes[node_id] = new_text
            self.current_revision += 1
            self.audit.append({"actor": actor, "revision": self.current_revision, "nodeId": node_id})
            return {"status": "committed", "revision": self.current_revision}


def run_patch_concurrency(output_dir: Path) -> dict[str, Any]:
    original = "现有研究尚未解释跨尺度因果关系。"
    store = RevisionStore({"node-1": original})
    base_hash = stable_hash(original)
    user_commit = store.commit(base_revision=1, node_id="node-1", expected_hash=base_hash, new_text="用户已补充前期实验边界。", actor="user")
    stale_ai_commit = store.commit(base_revision=1, node_id="node-1", expected_hash=base_hash, new_text="AI建议替换文本。", actor="ai_patch")

    race_store = RevisionStore({"node-1": original})
    race_results: list[dict[str, Any]] = []
    barrier = threading.Barrier(2)

    def contender(actor: str, value: str) -> None:
        barrier.wait()
        race_results.append(race_store.commit(base_revision=1, node_id="node-1", expected_hash=base_hash, new_text=value, actor=actor))

    threads = [
        threading.Thread(target=contender, args=("user", "用户并发修改")),
        threading.Thread(target=contender, args=("ai_patch", "AI并发修改")),
    ]
    for thread in threads:
        thread.start()
    for thread in threads:
        thread.join()

    report = {
        "spike": "patch-concurrency",
        "staleScenario": {
            "userCommit": user_commit,
            "staleAiCommit": stale_ai_commit,
            "finalText": store.nodes["node-1"],
            "staleOverwritePrevented": stale_ai_commit["status"] == "revision_conflict" and store.nodes["node-1"] == "用户已补充前期实验边界。",
        },
        "simultaneousScenario": {
            "results": race_results,
            "committedCount": sum(result["status"] == "committed" for result in race_results),
            "conflictCount": sum(result["status"] == "revision_conflict" for result in race_results),
        },
        "decision": "正式写入必须由 Revision Service 在同一临界区执行 baseRevision 比较与 currentRevision 推进；Patch 校验成功不赋予覆盖旧版本的权力。",
    }
    write_json(output_dir / "patch-concurrency" / "report.json", report)
    return report


class AuthorizationGateway:
    def __init__(self) -> None:
        self.authorization_revision = 1
        self.permissions = {"source-1": True}
        self.cache: dict[str, dict[str, Any]] = {}
        self.queue: list[dict[str, Any]] = []
        self.drafts: list[dict[str, Any]] = []
        self.audit: list[dict[str, Any]] = []

    def queue_call(self, task_id: str, source_ids: list[str]) -> None:
        self.queue.append({"taskId": task_id, "sourceIds": source_ids, "queuedAtAuthorizationRevision": self.authorization_revision})

    def cache_context(self, task_id: str, source_ids: list[str]) -> str:
        key = stable_hash({"taskId": task_id, "sourceIds": source_ids, "authorizationRevision": self.authorization_revision})
        self.cache[key] = {"sourceIds": source_ids, "authorizationRevision": self.authorization_revision}
        return key

    def create_draft(self, patch_id: str, source_ids: list[str], accepted: bool = False) -> None:
        self.drafts.append({"patchId": patch_id, "sourceIds": source_ids, "status": "accepted" if accepted else "draft"})

    def revoke(self, source_id: str) -> None:
        self.authorization_revision += 1
        self.permissions[source_id] = False
        for task in self.queue:
            if source_id in task["sourceIds"]:
                task["status"] = "authorization_revoked"
        self.cache = {key: value for key, value in self.cache.items() if source_id not in value["sourceIds"]}
        for draft in self.drafts:
            if source_id in draft["sourceIds"] and draft["status"] != "accepted":
                draft["status"] = "evidence_revoked"
        self.audit.append({"event": "evidence_revoked", "sourceId": source_id, "authorizationRevision": self.authorization_revision})

    def dispatch(self, task_id: str) -> dict[str, Any]:
        task = next(item for item in self.queue if item["taskId"] == task_id)
        denied = [source_id for source_id in task["sourceIds"] if not self.permissions.get(source_id, False)]
        if denied:
            return {"status": "blocked", "code": "evidence_authorization_revoked", "deniedSourceIds": denied}
        return {"status": "ready", "sourceIds": task["sourceIds"], "authorizationRevision": self.authorization_revision}


def run_authorization_propagation(output_dir: Path) -> dict[str, Any]:
    gateway = AuthorizationGateway()
    gateway.queue_call("task-queued", ["source-1"])
    cache_key = gateway.cache_context("task-cache", ["source-1"])
    gateway.create_draft("patch-draft", ["source-1"])
    gateway.create_draft("patch-accepted", ["source-1"], accepted=True)
    gateway.revoke("source-1")
    dispatch = gateway.dispatch("task-queued")
    report = {
        "spike": "authorization-propagation",
        "queuedCall": gateway.queue[0],
        "dispatchAfterRevocation": dispatch,
        "cachedContextRemoved": cache_key not in gateway.cache,
        "drafts": gateway.drafts,
        "acceptedRevisionRetainsAuditOnly": gateway.drafts[1]["status"] == "accepted" and len(gateway.audit) == 1,
        "futureModelUseBlocked": dispatch["status"] == "blocked",
        "decision": "队列仅保存 sourceId，不保存可直接发送的摘录；模型调用前必须重新查询当前授权。authorizationRevision 进入缓存键，撤权主动失效缓存和未接受 Patch，已接受 Revision 只保留不可执行的审计来源。",
    }
    write_json(output_dir / "authorization-propagation" / "report.json", report)
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Run isolated NSFC grant-platform technical spikes.")
    parser.add_argument("--output-dir", type=Path, default=Path(tempfile.gettempdir()) / "researchgpt-grant-spikes")
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    reports = {
        "docxRoundtrip": run_docx_roundtrip(args.output_dir),
        "anchorDrift": run_anchor_drift(args.output_dir),
        "patchConcurrency": run_patch_concurrency(args.output_dir),
        "authorizationPropagation": run_authorization_propagation(args.output_dir),
    }
    summary = {
        "outputDirectory": str(args.output_dir),
        "docxBodyStructurePreserved": reports["docxRoundtrip"]["bodyStructurePreserved"],
        "anchorSafeDecisionPassRate": reports["anchorDrift"]["safeDecisionPassRate"],
        "staleOverwritePrevented": reports["patchConcurrency"]["staleScenario"]["staleOverwritePrevented"],
        "simultaneousCommitArbitrated": (
            reports["patchConcurrency"]["simultaneousScenario"]["committedCount"] == 1
            and reports["patchConcurrency"]["simultaneousScenario"]["conflictCount"] == 1
        ),
        "authorizationPropagationPassed": all(
            (
                reports["authorizationPropagation"]["futureModelUseBlocked"],
                reports["authorizationPropagation"]["cachedContextRemoved"],
                reports["authorizationPropagation"]["drafts"][0]["status"] == "evidence_revoked",
            )
        ),
        "visualQa": "pending_external_render",
    }
    write_json(args.output_dir / "summary.json", summary)
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    required_checks = (
        summary["docxBodyStructurePreserved"],
        summary["anchorSafeDecisionPassRate"] == 1.0,
        summary["staleOverwritePrevented"],
        summary["simultaneousCommitArbitrated"],
        summary["authorizationPropagationPassed"],
    )
    if not all(required_checks):
        raise SystemExit(1)


if __name__ == "__main__":
    main()
